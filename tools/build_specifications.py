from __future__ import annotations

from pathlib import Path

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT / "Specifications techniques MultiPrep.docx"
PDF_PATH = ROOT / "Specifications techniques MultiPrep.pdf"
YELLOW = "F9AB00"
DARK = "202124"

SECTIONS = [
    (
        "1. Objet et périmètre",
        [
            "MultiPrep est une application Windows de préparation de dossiers PDF. Elle importe des PDF, documents Microsoft Word et images, permet leur réorganisation et produit un PDF final nommé selon les informations du dossier.",
            "Windows 10/11 · Python 3.10+ · PySide6 · livrable officiel : dist\\MultiPrep.exe.",
        ],
    ),
    (
        "2. Mode Gmail — mode par défaut",
        [
            "Glisser-déposer des pièces jointes Gmail dans la grande zone intégrée.",
            "Gestion native des fichiers virtuels et du transfert asynchrone Windows.",
            "Copie puis collage des images présentes dans le corps du mail Gmail.",
            "Collage de fichiers JPG/PNG copiés depuis l’Explorateur.",
            "Imports successifs sans redémarrage ni fenêtre auxiliaire visible.",
        ],
    ),
    (
        "3. Mode classique",
        [
            "Import depuis l’Explorateur, le Bureau ou une application de bureau exposant un vrai fichier, dont Outlook bureau.",
            "Outlook est une source compatible parmi d’autres ; le mode n’en dépend pas.",
            "Collage de captures et d’images locales.",
        ],
    ),
    (
        "4. Formats et conversions",
        [
            "PDF : .pdf · Microsoft Word : .doc/.docx · Images : .jpg/.jpeg/.png.",
            "Microsoft Word est requis localement pour convertir DOC/DOCX en PDF.",
            "Les images sont encapsulées dans un PDF avant assemblage.",
        ],
    ),
    (
        "5. Architecture",
        [
            "main_window.py : fenêtre, thèmes et cycle de vie Gmail.",
            "editor_view.py : formulaire, en-tête et consignes.",
            "page_grid_list.py et page_card_delegate.py : grille rapide, sélection et rendu léger.",
            "drop_service.py : fichiers locaux, virtuels, images et HTML.",
            "gmail_import_service.py et GmailDropHelper.cs : intégration et transfert natif Windows.",
            "pdf_service.py, thumbnail_service.py et word_service.py : PDF, aperçus et conversion Word.",
        ],
    ),
    (
        "6. Performances",
        [
            "Modèles créés sans rendu bloquant et cartes dessinées par un delegate léger.",
            "Jusqu’à quatre lots de miniatures en parallèle, ouverture groupée des PDF et cache local.",
            "Ajout incrémental sans reconstruction des pages existantes.",
            "Référence : 500 pages injectées en 10–12 ms et 120 miniatures en environ 0,63 seconde sur le poste de développement.",
        ],
    ),
    (
        "7. Identité visuelle 2.0.0",
        [
            "Mode Gmail : blanc, jaune Google #F9AB00 et texte #202124.",
            "Mode classique : identité bleue historique conservée.",
            "Logo 2.0.0 blanc/jaune dans l’exécutable et les documents.",
        ],
    ),
    (
        "8. Sécurité, validation et distribution",
        [
            "Aucun identifiant Gmail n’est lu ou stocké ; seules les données remises par Windows sont utilisées.",
            "Le cache temporaire est nettoyé au lancement et à la fermeture normale.",
            "Validation : compileall, unittest et git diff --check.",
            "Build PyInstaller via MultiPrep.spec avec un workpath temporaire.",
        ],
    ),
]


def build_docx() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Cm(1.5)
    section.left_margin = section.right_margin = Cm(1.7)

    normal = document.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(10.5)
    for style_name in ("Title", "Heading 1"):
        document.styles[style_name].font.name = "Segoe UI"

    title = document.add_heading("MultiPrep", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor.from_string(DARK)
    title.runs[0].font.size = Pt(30)
    subtitle = document.add_paragraph("SPÉCIFICATIONS TECHNIQUES — VERSION 2.0.0")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(15)
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(YELLOW)
    meta = document.add_paragraph("Édition du 3 juillet 2026 · Google Workspace et applications de bureau")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for heading, items in SECTIONS:
        paragraph = document.add_heading(heading, level=1)
        paragraph.runs[0].font.color.rgb = RGBColor.from_string(DARK)
        for item in items:
            document.add_paragraph(item, style="List Bullet")

    footer = section.footer.paragraphs[0]
    footer.text = "MultiPrep 2.0.0 · Spécifications techniques · 3 juillet 2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(DOCX_PATH)


def build_pdf() -> None:
    body = [
        "<h1>MultiPrep</h1>",
        "<div class='subtitle'>SPÉCIFICATIONS TECHNIQUES — VERSION 2.0.0</div>",
        "<div class='meta'>Édition du 3 juillet 2026 · Google Workspace et applications de bureau</div>",
    ]
    for heading, items in SECTIONS:
        body.append(f"<h2>{heading}</h2><ul>")
        body.extend(f"<li>{item}</li>" for item in items)
        body.append("</ul>")
    html = "<html><body>" + "".join(body) + "</body></html>"
    css = f"""
        body {{ font-family: sans-serif; color: #{DARK}; font-size: 10.5pt; line-height: 1.45; }}
        h1 {{ text-align: center; font-size: 30pt; margin: 8pt 0 0; border-bottom: 5px solid #{YELLOW}; }}
        .subtitle {{ text-align: center; color: #9a5d00; font-size: 15pt; font-weight: bold; margin: 8pt; }}
        .meta {{ text-align: center; color: #5f6368; margin-bottom: 20pt; }}
        h2 {{ font-size: 16pt; border-left: 6px solid #{YELLOW}; padding-left: 8pt; margin-top: 18pt; }}
        li {{ margin-bottom: 4pt; }}
    """
    story = fitz.Story(html=html, user_css=css)
    writer = fitz.DocumentWriter(str(PDF_PATH))
    page_rect = fitz.paper_rect("a4")
    content_rect = page_rect + (42, 42, -42, -42)
    more = True
    while more:
        device = writer.begin_page(page_rect)
        more, _ = story.place(content_rect)
        story.draw(device)
        writer.end_page()
    writer.close()


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(f"Generated: {DOCX_PATH.name}, {PDF_PATH.name}")
